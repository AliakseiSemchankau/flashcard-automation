from docx import Document
from docx.table import Table
from docx.shared import Pt


def write_to_table(table: Table, sentences: list[str], bold_words: list[str], font: str, font_size: int) -> None:
    """
    Writes sentences into a docx table, applying bold formatting to specific words,
    and sets the font and font size.
    """
    for index, (sentence, bold_word) in enumerate(zip(sentences, bold_words)):
        # Calculate cell position
        row_index = index // len(table.columns)
        col_index = index % len(table.columns)

        cell = table.rows[row_index].cells[col_index]
        cell.text = ""  # Clear existing content
        paragraph = cell.add_paragraph()

        # Add text with bold formatting
        before, sep, after = sentence.partition(bold_word)
        if before:
            paragraph.add_run(before)
        if sep:
            bold_run = paragraph.add_run(sep)
            bold_run.bold = True
        if after:
            paragraph.add_run(after)

        # Set font properties
        for run in paragraph.runs:
            run.font.name = font
            run.font.size = Pt(font_size)


def reverse(sequence: list[str], n_cols: int, table_size: int) -> list[str]:
    """
    Reverses each row of a sequence when arranged in a table format.
    """
    new_sequence = [" "] * table_size
    for index, value in enumerate(sequence):
        i = index // n_cols
        j = index % n_cols
        new_sequence[i * n_cols + n_cols - 1 - j] = value
    return new_sequence



class DocxClient:
    """
    A client for managing Word documents for flashcards.
    """

    def __init__(
    self,
    template_path: str = "template.docx",
    folder: str = "output/",
    prefix: str = "flashcards",
    n_rows: int = 4,
    n_cols: int = 3,
    base_font: str = "Arial",
    target_font: str = "Arial",
    font_size: int = 15
    ):
        """
        Initializes the DocxClient with default settings for folder, template, prefix,
        table dimensions, and font size.
        """
        self.template = template_path
        self.folder = folder
        self.flashcards_prefix = prefix
        self.n_rows = n_rows
        self.n_cols = n_cols
        self.table_size = n_rows * n_cols
        self.base_font = base_font
        self.target_font = target_font
        self.font_size = font_size


    def _name_to_path(self, filename: str) -> str:
        """
        Constructs the full file path.
        """
        return f"{self.folder}{self.flashcards_prefix}{filename}.docx"


    def _write_records_to_docx(self, filename: str, records: dict) -> None:
        """
        Writes records into a Word document template and saves it.
        """
        doc = Document(self.template)

        # Validate tables
        if len(doc.tables) < 2:
            print(f"Error: Template must contain at least two tables. Found {len(doc.tables)}.")
            return

        for i in range(2):
            table_dims = (len(doc.tables[i].rows), len(doc.tables[i].columns))
            expected_dims = (self.n_rows, self.n_cols)
            if table_dims != expected_dims:
                print(f"Error: Table #{i + 1} dimensions {table_dims} do not match expected {expected_dims}.")
                return

        # Write base and target language sentences
        # When writing the latter ones, reverse the rows
        write_to_table(
            doc.tables[0],
            records["base_sentences"],
            records["base_bold_words"],
            self.base_font,
            self.font_size,
        )
        write_to_table(
            doc.tables[1],
            reverse(records["target_sentences"], self.n_cols, self.table_size),
            reverse(records["target_bold_words"], self.n_cols, self.table_size),
            self.target_font,
            self.font_size,
        )

        # Save document
        file_path = self._name_to_path(filename)
        doc.save(file_path)
        print(f"Document saved to: {file_path}")


    def write_records(self, filename: str, records: dict) -> list[tuple[str, str]]:
        """
        Writes records into one or more Word documents and returns file paths and names.
        """
        if records["n_records"] <= self.table_size:
            self._write_records_to_docx(filename, records)
            return [(self._name_to_path(filename), filename)]

        n_files = (records["n_records"] + self.table_size - 1) // self.table_size
        filenames = []

        for i in range(n_files):
            suffix = f"-{i + 1}"
            chunk = slice(i * self.table_size, min((i + 1) * self.table_size, records["n_records"]))
            filenames.append(filename + suffix)
            self._write_records_to_docx(
                filenames[-1],
                {key: val[chunk] if isinstance(val, list) else val for key, val in records.items()},
            )

        return [(self._name_to_path(name), self.flashcards_prefix + name) for name in filenames]



